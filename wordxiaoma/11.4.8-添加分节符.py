# -*- coding: utf-8 -*-
"""
Created on Sat Feb 19 18:58:15 2022

@author: 小码哥
"""
# 导入 python-docx 库
from docx import Document
# 导入Inches函数
from docx.shared import Inches

# 用 list-list 存储表格数据
data = [['省市','2021GDP（亿元）','2020GDP（亿元）','2019GDP（亿元）'],
        ['广东省','123469.67','110760.94','107671.07'],
        ['江苏省','116364.2','102719','99631.52'],
        ['山东省','83095.9','73129','62352']]

# 创建空白的 Word 文档
myDoc = Document()

# 添加标题:
myDoc.add_heading('近三年省市 GDP 前三名数据统计表')

############
# 添加表格
############
# 行列数
rows = 4
cols = 4

# 添加表格的方法 add_table(rows, cols)
addedTable = myDoc.add_table(rows, cols)

# 为表格添加数据
for row in range(rows):
    for col in range(cols):
        addedTable.cell(row,col).text = str(data[row][col])

'''
添加分节符
4 - 分节符(奇数页)
3 - 分节符(偶数页)
2 - 分节符(下一页)
1 - 分节符(节的结尾)
0 - 分节符(连续)
'''
myDoc.add_section(4)

'''
添加图片
'''
# 图片路径
pic_path = 'picture/封面.jpg'
# 添加图片的方法add_picture(),并且指定宽度 width
myDoc.add_picture(pic_path, width = Inches(6))

# 保存文档
myDoc.save('generated_docx/11.4.8-添加分节符.docx')

# 提示信息
print('添加成功！')
