# -*- coding: utf-8 -*-
"""
Created on Thu Feb 17 00:24:25 2022

@author: 小码哥
"""
# 导入 python-docx 库
from docx import Document

# 创建空白的 Word 文档
myDoc = Document()

# 添加标题:
myDoc.add_heading('公众号“七天小码哥”粉丝分析')
# 默认在末尾 添加段落
# 添加第一个段落
para_0 = myDoc.add_paragraph("用户增长是一个公众号的命脉所在。")

# 添加第二个段落
para_1 = myDoc.add_paragraph("用户总数、用户流失、用户新增")

# 在第二个 段落 para_1 之前插入一个新的段落。到此文章共有3个段落。
para_1.insert_paragraph_before("它分为几个类别：")

# 保存文档
myDoc.save('generated_docx/11.4.3-添加多个段落.docx')

print('创建成功！')