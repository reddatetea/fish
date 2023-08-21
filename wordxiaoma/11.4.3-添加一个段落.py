# -*- coding: utf-8 -*-
"""
Created on Thu Feb 17 00:24:25 2022

@author: 小码哥
"""
# 导入 python-docx 库
from docx import Document

# 创建空白的 Word 文档
myDoc = Document()

# 添加标题:
myDoc.add_heading('公众号“七天小码哥”粉丝分析')
# 默认在末尾 添加段落
# 添加第一个段落
para_0 = myDoc.add_paragraph("用户增长是一个公众号的命脉所在。")

# 保存文档
myDoc.save('generated_docx/11.4.3-添加一个段落.docx')

print('创建成功！')


