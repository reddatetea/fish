# -*- coding: utf-8 -*-
"""
Created on Sat Feb 19 22:26:22 2022

@author: 小码哥
"""
from docx import Document
import pandas as pd

stu = pd.read_excel('data/学生成绩.xlsx')
# 排序
stu.sort_values(by='分数', inplace=True, ascending=False)
stu.reset_index(drop=True, inplace=True)
first_student = stu['姓名'][0]
first_score = stu['分数'][0]
number_student = len(stu['姓名'])

doc = Document()

doc.add_heading('一班学生期末考试情况', level = 0)

# 添加第一个段落
p = doc.add_paragraph('分数最高的学生是')

p.add_run(str(first_student)).bold = True
p.add_run('，并且分数是')
p.add_run(str(first_score)).bold=True
p.add_run('分。')

# 添加第二个段落
p1 = doc.add_paragraph('总共有{}名学生参加了考试，考试总体情况是：'.format(number_student))

# 添加表格
table = doc.add_table(number_student+1, 2)
table.cell(0,0).text = '学生成绩'
table.cell(0,1).text = '学生分数'
table.style = 'Medium Grid 1 Accent 4'

# 添加数据
for i,(inde,row) in enumerate(stu.iterrows()):
    table.cell(i+1,0).text=str(row['姓名'])
    table.cell(i+1,1).text=str(row['分数'])

# 添加图片
doc.add_picture('picture/学生成绩.png')
doc.save('generated_docx/11.5-案例.docx')